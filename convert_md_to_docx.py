import markdown2
from docx import Document
import sys
import os

# Archivo markdown de entrada y docx de salida
def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    html = markdown2.markdown(md_content)
    
    # Crear documento Word
    doc = Document()
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    for elem in soup.descendants:
        if elem.name == 'h1':
            doc.add_heading(elem.text, level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.text, level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.text, level=3)
        elif elem.name == 'li':
            doc.add_paragraph(elem.text, style='List Bullet')
        elif elem.name == 'p':
            doc.add_paragraph(elem.text)
        elif elem.name == 'pre':
            doc.add_paragraph(elem.text, style='Intense Quote')
    doc.save(docx_path)

if __name__ == "__main__":
    md_file = sys.argv[1] if len(sys.argv) > 1 else 'INFORME_PROYECTO_GEOAPP.md'
    docx_file = os.path.splitext(md_file)[0] + '.docx'
    convert_md_to_docx(md_file, docx_file)
    print(f"Archivo Word generado: {docx_file}")
